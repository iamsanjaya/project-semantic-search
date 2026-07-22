from docx import Document
from docx.table import Table
from typing import List
from src.ingestion.base_ingestor import BaseIngestor
from src.domain.page import Page


class DOCXIngestor(BaseIngestor):
    """
    Production-safe DOCX ingestor.

    Improves retrieval by grouping paragraphs into heading-scoped sections
    (instead of one Page per paragraph), preserving heading context, and
    extracting tables as their own structured Pages.

    Design decisions (confirmed against real SajiloMart doc structure):
    - Recognizes standard Word heading styles: "Heading 1/2/3".
    - Skips "toc *" styles entirely — these duplicate real heading text
      and would otherwise double-ingest every section title.
    - Guards against mis-styled paragraphs (a heading-styled paragraph
      that is actually long, sentence-like body text) by falling back
      to treating it as body content instead of starting a new section.
    - Caps section size; overly long sections are split into multiple
      Pages so the chunker never receives an oversized block.
    - Tables are extracted separately, row-wise (similar pattern to
      CSVIngestor), and tagged with source="docx_table" for traceability.
    """

    HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3"}
    MAX_SECTION_CHARS = 1500
    # A "heading" longer than this is almost certainly a mis-styled
    # body paragraph, not a real section title.
    MAX_HEADING_CHARS = 100

    def _load(self, file_path: str) -> List[Page]:
        document = Document(file_path)
        file_name = file_path.split("/")[-1]

        pages: List[Page] = []
        page_number = 1

        current_heading_path: List[str] = (
            []
        )  # e.g. ["6. Database Design", "6.1 Core Entities & Fields"]
        current_heading_levels: List[int] = (
            []
        )  # parallel stack of heading levels (1/2/3)
        current_body_parts: List[str] = []

        def flush_section():
            """Turn accumulated body text under the current heading into one or more Pages."""
            nonlocal page_number
            if not current_body_parts:
                return

            heading_label = (
                " > ".join(current_heading_path)
                if current_heading_path
                else "Untitled Section"
            )
            body_text = " ".join(current_body_parts).strip()
            if not body_text:
                return

            # Split into size-capped chunks if the section is too long
            segments = self._split_by_size(body_text, self.MAX_SECTION_CHARS)

            for seg_index, segment in enumerate(segments):
                # Prepend heading context so retrieval keeps section identity
                text_with_context = (
                    f"{heading_label}: {segment}" if heading_label else segment
                )

                pages.append(
                    Page(
                        file_name=file_name,
                        page_number=page_number,
                        text=text_with_context,
                        metadata={
                            "source": "docx",
                            "file_path": file_path,
                            "heading_path": heading_label,
                            "heading_level": (
                                current_heading_levels[-1]
                                if current_heading_levels
                                else None
                            ),
                            "section_segment_index": seg_index,
                            "char_count": len(text_with_context),
                        },
                    )
                )
                page_number += 1

        for paragraph in document.paragraphs:
            style_name = (paragraph.style.name if paragraph.style else None) or "Normal"
            text = paragraph.text.strip()

            if not text:
                continue

            # Skip auto-generated Table of Contents entries entirely —
            # they duplicate real heading text and would cause double-ingestion.
            if style_name.lower().startswith("toc"):
                continue

            is_heading_style = style_name in self.HEADING_STYLES
            is_plausible_heading = (
                is_heading_style and len(text) <= self.MAX_HEADING_CHARS
            )

            if is_heading_style and not is_plausible_heading:
                # Mis-styled paragraph (styled as heading but reads like body text).
                # Treat as body content instead of starting a new section.
                current_body_parts.append(text)
                continue

            if is_plausible_heading:
                # New heading encountered — flush whatever section we were building.
                flush_section()
                current_body_parts = []

                level = int(style_name[-1])  # "Heading 1" -> 1, "Heading 2" -> 2, etc.

                # Pop the heading stack back to the correct depth for this level
                while current_heading_levels and current_heading_levels[-1] >= level:
                    current_heading_levels.pop()
                    current_heading_path.pop()

                current_heading_levels.append(level)
                current_heading_path.append(text)
                continue

            # Regular body text (includes "Normal" and "List Paragraph" styles)
            current_body_parts.append(text)

        # Flush whatever section remains at end of document
        flush_section()

        # -----------------------------
        # TABLE EXTRACTION
        # -----------------------------
        for table_index, table in enumerate(document.tables):
            table_pages = self._extract_table(
                table, file_name, file_path, table_index, page_number
            )
            pages.extend(table_pages)
            page_number += len(table_pages)

        return pages

    # -----------------------------
    # HELPERS
    # -----------------------------
    def _split_by_size(self, text: str, max_chars: int) -> List[str]:
        """Split long section text into size-capped segments on sentence-ish boundaries."""
        if len(text) <= max_chars:
            return [text]

        segments = []
        remaining = text
        while len(remaining) > max_chars:
            # Try to split at the last sentence boundary before max_chars
            split_point = remaining.rfind(". ", 0, max_chars)
            if split_point == -1:
                split_point = max_chars  # hard cut if no natural boundary found
            segments.append(remaining[: split_point + 1].strip())
            remaining = remaining[split_point + 1 :].strip()
        if remaining:
            segments.append(remaining)
        return segments

    def _extract_table(
        self,
        table: Table,
        file_name: str,
        file_path: str,
        table_index: int,
        start_page_number: int,
    ) -> List[Page]:
        """Convert a docx table into row-wise text Pages, similar to CSVIngestor's pattern."""
        rows = table.rows
        if not rows:
            return []

        header_cells = [cell.text.strip() for cell in rows[0].cells]
        pages: List[Page] = []
        page_number = start_page_number

        for row_index, row in enumerate(rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            row_parts = []
            for col_name, value in zip(header_cells, cells):
                if not value:
                    continue
                row_parts.append(f"{col_name}: {value}")

            row_text = ". ".join(row_parts).strip()
            if not row_text or len(row_text) < 10:
                continue

            pages.append(
                Page(
                    file_name=file_name,
                    page_number=page_number,
                    text=row_text,
                    metadata={
                        "source": "docx_table",
                        "file_path": file_path,
                        "table_index": table_index,
                        "row_index": row_index,
                        "columns": header_cells,
                        "char_count": len(row_text),
                    },
                )
            )
            page_number += 1

        return pages
